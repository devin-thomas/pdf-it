"""File and transcript ingestion helpers for mixed-source document creation."""

from __future__ import annotations

import io
import json
import mimetypes
import re
from collections.abc import Callable, Sequence
from dataclasses import dataclass
from importlib import import_module
from urllib.error import HTTPError, URLError
from urllib.parse import parse_qs, urlparse
from urllib.request import Request, urlopen

import pandas as pd
from google import genai
from google.genai import types as genai_types
from openai import OpenAI
from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api._errors import (
    NoTranscriptFound,
    TranscriptsDisabled,
    VideoUnavailable,
    YouTubeTranscriptApiException,
)

from .config import (
    MAX_DATAFRAME_COLUMNS,
    MAX_DATAFRAME_ROWS,
    MAX_SOURCE_CHARACTERS,
    MAX_SOURCE_UPLOADS,
    MAX_UPLOAD_BYTES,
    MAX_YOUTUBE_LINKS,
    PROVIDER_CONFIGS,
    Provider,
)
from .providers import safe_provider_error
from .validation import InputValidationError, decode_text_upload

ProgressReporter = Callable[[str, int], None]

TEXT_EXTENSIONS = {".txt", ".md", ".markdown", ".mdx"}
DOCUMENT_EXTENSIONS = {".pdf", ".doc", ".docx"}
SPREADSHEET_EXTENSIONS = {".csv", ".xlsx"}
AUDIO_EXTENSIONS = {".aac", ".flac", ".m4a", ".mp3", ".ogg", ".oga", ".wav", ".webm"}
VIDEO_EXTENSIONS = {".avi", ".m4v", ".mkv", ".mov", ".mp4", ".mpeg", ".mpg", ".wmv"}

CODE_LANGUAGES = {
    ".py": ("Python", "python"),
    ".js": ("JavaScript", "javascript"),
    ".jsx": ("JavaScript", "jsx"),
    ".ts": ("TypeScript", "typescript"),
    ".tsx": ("TypeScript", "tsx"),
    ".java": ("Java", "java"),
    ".c": ("C", "c"),
    ".h": ("C", "c"),
    ".cpp": ("C++", "cpp"),
    ".cc": ("C++", "cpp"),
    ".cxx": ("C++", "cpp"),
    ".hpp": ("C++", "cpp"),
    ".cs": ("C#", "csharp"),
    ".php": ("PHP", "php"),
    ".go": ("Go", "go"),
    ".rs": ("Rust", "rust"),
    ".swift": ("Swift", "swift"),
    ".kt": ("Kotlin", "kotlin"),
    ".kts": ("Kotlin", "kotlin"),
    ".rb": ("Ruby", "ruby"),
    ".sql": ("SQL", "sql"),
    ".html": ("HTML", "html"),
    ".css": ("CSS", "css"),
}

UNSUPPORTED_CODE_EXTENSIONS = {
    ".bat",
    ".dart",
    ".elm",
    ".erl",
    ".ex",
    ".lua",
    ".m",
    ".pl",
    ".ps1",
    ".r",
    ".scala",
    ".scss",
    ".sh",
    ".vb",
}

YOUTUBE_VIDEO_ID_PATTERN = re.compile(r"^[A-Za-z0-9_-]{11}$")
YOUTUBE_SHORT_HOSTS = {"youtu.be", "www.youtu.be"}
YOUTUBE_HOST_SUFFIXES = ("youtube.com", "youtube-nocookie.com")


@dataclass(frozen=True)
class SourceUpload:
    name: str
    data: bytes


@dataclass(frozen=True)
class SourceSegment:
    label: str
    text: str


@dataclass(frozen=True)
class PreparedSources:
    source: str
    source_count: int
    used_local_plan: bool = False


def prepare_source_text(
    typed_text: str,
    uploads: Sequence[SourceUpload],
    youtube_urls: Sequence[str],
    provider: Provider,
    api_key: str,
    model_name: str,
    *,
    progress: ProgressReporter | None = None,
) -> PreparedSources:
    """Combine typed, uploaded, and transcript-based sources into one bounded prompt."""
    cleaned_urls = tuple(url.strip() for url in youtube_urls if url.strip())
    if len(uploads) > MAX_SOURCE_UPLOADS:
        raise InputValidationError(f"Add at most {MAX_SOURCE_UPLOADS} uploaded files at once.")
    if len(cleaned_urls) > MAX_YOUTUBE_LINKS:
        raise InputValidationError(
            f"Add at most {MAX_YOUTUBE_LINKS} YouTube links or video IDs at once."
        )

    segments: list[SourceSegment] = []
    if typed_text.strip():
        segments.append(SourceSegment("Free text", typed_text.strip()))

    total_uploads = len(uploads) + len(cleaned_urls)
    for index, upload in enumerate(uploads, start=1):
        if progress is not None:
            progress(f"Reading file {index} of {total_uploads}: {upload.name}", 12 + index * 6)
        segments.append(
            SourceSegment(
                label=upload.name,
                text=extract_upload_text(upload, provider, api_key, model_name),
            )
        )

    for offset, url in enumerate(cleaned_urls, start=1):
        if progress is not None:
            progress(
                f"Pulling YouTube transcript {offset} of {len(cleaned_urls)}",
                46 + offset * 6,
            )
        segments.append(
            SourceSegment(
                label=f"YouTube transcript {offset}",
                text=fetch_youtube_transcript(url),
            )
        )

    if not segments:
        raise InputValidationError(
            "Add source text, upload a supported file, or include a YouTube link."
        )

    return PreparedSources(
        source=combine_source_segments(segments),
        source_count=len(segments),
    )


def extract_upload_text(
    upload: SourceUpload,
    provider: Provider,
    api_key: str,
    model_name: str,
) -> str:
    """Extract user-readable text from a supported upload."""
    if len(upload.data) > MAX_UPLOAD_BYTES:
        raise InputValidationError("Each uploaded file must be 10 MB or smaller.")

    extension = normalized_extension(upload.name)
    if extension in TEXT_EXTENSIONS:
        return decode_text_upload(upload.data)
    if extension in CODE_LANGUAGES:
        language, fence = CODE_LANGUAGES[extension]
        code = decode_text_upload(upload.data)
        return f"Language: {language}\nFile: {upload.name}\n```{fence}\n{code}\n```"
    if extension in SPREADSHEET_EXTENSIONS:
        return extract_tabular_text(upload.name, upload.data, extension)
    if extension in DOCUMENT_EXTENSIONS:
        return extract_document_text(upload.name, upload.data)
    if extension == ".gdoc":
        return extract_gdoc_text(upload.data)
    if extension in AUDIO_EXTENSIONS:
        return transcribe_audio(upload, provider, api_key, model_name)
    raise InputValidationError(unsupported_extension_message(extension))


def combine_source_segments(segments: Sequence[SourceSegment]) -> str:
    """Render source segments with labels and fit them within the shared prompt budget."""
    if not segments:
        raise InputValidationError("Add source text or a supported upload.")

    formatted = [render_source_segment(segment.label, segment.text) for segment in segments]
    source = "\n\n".join(formatted)
    if len(source) <= MAX_SOURCE_CHARACTERS:
        return source

    headers = [render_source_segment(segment.label, "") for segment in segments]
    content_budget = MAX_SOURCE_CHARACTERS - sum(len(header) for header in headers) - (
        2 * (len(segments) - 1)
    )
    if content_budget <= 600:
        raise InputValidationError(
            f"Source text must fit within {MAX_SOURCE_CHARACTERS:,} characters."
        )

    share = max(500, content_budget // len(segments))
    budgets = [min(len(segment.text), share) for segment in segments]
    leftover = content_budget - sum(budgets)
    while leftover > 0:
        updated = False
        for index, segment in enumerate(segments):
            remaining = len(segment.text) - budgets[index]
            if remaining <= 0:
                continue
            step = min(remaining, leftover)
            budgets[index] += step
            leftover -= step
            updated = True
            if leftover == 0:
                break
        if not updated:
            break

    trimmed = [
        render_source_segment(segment.label, truncate_text(segment.text, budgets[index]))
        for index, segment in enumerate(segments)
    ]
    final_source = "\n\n".join(trimmed)
    return final_source[:MAX_SOURCE_CHARACTERS].rstrip()


def extract_document_text(name: str, data: bytes) -> str:
    """Extract document text without crashing when optional parsers are absent."""
    extension = normalized_extension(name)
    if extension == ".pdf":
        return extract_pdf_text(name, data)
    if extension == ".docx":
        return extract_docx_text(name, data)
    if extension == ".doc":
        return extract_legacy_doc_text(name, data)

    raise InputValidationError(
        f"`{name}` is not a supported document upload yet. "
        "Try exporting it as `.pdf` or `.docx`."
    )


def extract_pdf_text(name: str, data: bytes) -> str:
    """Read searchable PDF text with pypdf."""
    try:
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # pragma: no cover - parser-specific failure details vary.
        raise InputValidationError(
            f"pdf-it could not read `{name}` as a PDF. Try exporting it again and re-uploading."
        ) from exc

    blocks = []
    for page in reader.pages:
        text = (page.extract_text() or "").strip()
        if text:
            blocks.append(text)
    if not blocks:
        raise InputValidationError(f"`{name}` did not contain extractable text.")
    return "\n\n".join(blocks)


def extract_docx_text(name: str, data: bytes) -> str:
    """Read DOCX text using python-docx."""
    try:
        from docx import Document

        document = Document(io.BytesIO(data))
    except Exception as exc:  # pragma: no cover - parser-specific failure details vary.
        raise InputValidationError(
            f"pdf-it could not read `{name}` as a Word document. "
            "Try exporting it again and re-uploading."
        ) from exc

    blocks = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                blocks.append(" | ".join(cells))
    if not blocks:
        raise InputValidationError(f"`{name}` did not contain extractable text.")
    return "\n\n".join(blocks)


def load_udoc():
    """Import the optional legacy document parser when it is available."""
    try:
        return import_module("udoc")
    except ModuleNotFoundError:
        return None


def extract_legacy_doc_text(name: str, data: bytes) -> str:
    """Read older `.doc` files only when the optional parser is present."""
    udoc = load_udoc()
    if udoc is None:
        raise InputValidationError(
            f"`{name}` is an older `.doc` file, which is not available in this deployment. "
            "Save it as `.docx` or `.pdf` and upload it again."
        )

    try:
        document = udoc.extract_bytes(data)
    except Exception as exc:  # pragma: no cover - legacy parser failures vary by file.
        raise InputValidationError(
            f"pdf-it could not read `{name}`. Try exporting it again and re-uploading."
        ) from exc

    blocks = [block.text.strip() for block in document.blocks() if block.text.strip()]
    if not blocks:
        raise InputValidationError(f"`{name}` did not contain extractable text.")
    return "\n\n".join(blocks)


def extract_gdoc_text(data: bytes) -> str:
    """Fetch text from a Google Docs pointer file when the linked doc is accessible."""
    try:
        payload = json.loads(data.decode("utf-8-sig"))
    except (UnicodeDecodeError, json.JSONDecodeError) as exc:
        raise InputValidationError(
            "The uploaded `.gdoc` file could not be read. Download the Google Doc as "
            "`.docx` or paste the text instead."
        ) from exc

    candidate_url = payload.get("url") or payload.get("doc_url") or ""
    doc_id = payload.get("doc_id") or extract_google_doc_id(candidate_url)
    if not doc_id:
        raise InputValidationError(
            "The `.gdoc` file did not include a usable Google Docs link. Download the "
            "doc as `.docx` or paste the text instead."
        )

    export_url = f"https://docs.google.com/document/d/{doc_id}/export?format=txt"
    request = Request(export_url, headers={"User-Agent": "pdf-it/0.1"})
    try:
        with urlopen(request, timeout=30) as response:
            body = response.read()
    except HTTPError as exc:
        if exc.code in {401, 403}:
            raise InputValidationError(
                "That Google Doc could not be exported from this app. If the doc is "
                "private, download it as `.docx` or paste the text instead."
            ) from exc
        raise InputValidationError(
            "pdf-it could not export that Google Doc right now."
        ) from exc
    except URLError as exc:
        raise InputValidationError(
            "pdf-it could not reach Google Docs to export that `.gdoc` file."
        ) from exc

    try:
        text = body.decode("utf-8-sig").strip()
    except UnicodeDecodeError as exc:
        raise InputValidationError(
            "Google Docs returned unreadable text for that `.gdoc` file."
        ) from exc
    if not text:
        raise InputValidationError(
            "That Google Doc exported successfully but did not contain readable text."
        )
    return text


def extract_tabular_text(name: str, data: bytes, extension: str) -> str:
    """Load CSV/XLSX data into pandas before converting it into prompt text."""
    try:
        if extension == ".csv":
            frame = pd.read_csv(io.BytesIO(data))
            return render_dataframe(frame, f"Sheet: {name}")

        workbook = pd.read_excel(io.BytesIO(data), sheet_name=None)
    except Exception as exc:  # pragma: no cover - parser-specific failure details vary.
        raise InputValidationError(
            f"pdf-it could not read `{name}` as a spreadsheet. Check that the file "
            "opens correctly and try again."
        ) from exc

    if not workbook:
        raise InputValidationError(f"`{name}` did not contain any readable sheets.")

    rendered_sheets = [
        render_dataframe(frame, f"Sheet: {sheet_name}")
        for sheet_name, frame in workbook.items()
    ]
    return "\n\n".join(rendered_sheets)


def render_dataframe(frame: pd.DataFrame, heading: str) -> str:
    """Convert a dataframe to compact plain text while preserving structure."""
    safe_frame = frame.copy()
    safe_frame.columns = [str(column) for column in safe_frame.columns]
    preview = safe_frame.fillna("")
    row_count, column_count = preview.shape
    clipped = preview.iloc[:MAX_DATAFRAME_ROWS, :MAX_DATAFRAME_COLUMNS]
    rendered = clipped.to_string(index=False, max_rows=MAX_DATAFRAME_ROWS)
    notes: list[str] = [f"{heading}", f"Rows: {row_count}, Columns: {column_count}"]
    if row_count > MAX_DATAFRAME_ROWS or column_count > MAX_DATAFRAME_COLUMNS:
        notes.append(
            f"Preview limited to the first {MAX_DATAFRAME_ROWS} rows and "
            f"{MAX_DATAFRAME_COLUMNS} columns."
        )
    notes.append(rendered)
    return "\n".join(notes)


def transcribe_audio(
    upload: SourceUpload,
    provider: Provider,
    api_key: str,
    model_name: str,
) -> str:
    """Transcribe audio uploads using the currently selected provider when possible."""
    mime_type = mimetypes.guess_type(upload.name)[0] or "application/octet-stream"
    if provider is Provider.GEMINI:
        try:
            client = genai.Client(api_key=api_key.strip())
            response = client.models.generate_content(
                model=model_name,
                contents=[
                    "Produce a clean transcript of this audio. Preserve speaker wording, "
                    "omit filler only when it is clearly non-semantic noise, and do not "
                    "summarize.",
                    genai_types.Part.from_bytes(data=upload.data, mime_type=mime_type),
                ],
            )
        except Exception as exc:  # pragma: no cover - provider-specific transport details vary.
            raise safe_provider_error(exc) from exc
        transcript = (response.text or "").strip()
        if transcript:
            return transcript
        raise InputValidationError(f"Gemini returned an empty transcript for `{upload.name}`.")

    if provider is Provider.OPENAI:
        try:
            client = OpenAI(api_key=api_key.strip())
            buffer = io.BytesIO(upload.data)
            buffer.name = upload.name
            response = client.audio.transcriptions.create(
                file=buffer,
                model=PROVIDER_CONFIGS[provider].transcription_model,
                response_format="text",
            )
        except Exception as exc:  # pragma: no cover - provider-specific transport details vary.
            raise safe_provider_error(exc) from exc
        transcript = response if isinstance(response, str) else getattr(response, "text", "")
        transcript = transcript.strip()
        if transcript:
            return transcript
        raise InputValidationError(f"OpenAI returned an empty transcript for `{upload.name}`.")

    raise InputValidationError(
        "Audio transcription is currently available with Gemini or OpenAI. "
        "Switch providers or upload a transcript."
    )


def fetch_youtube_transcript(video_input: str) -> str:
    """Fetch a YouTube transcript when public captions are available."""
    video_id = extract_youtube_video_id(video_input)
    if not video_id:
        raise InputValidationError(
            "Enter a valid YouTube link or 11-character video ID to import a transcript."
        )

    api = YouTubeTranscriptApi()
    try:
        fetched = api.fetch(video_id, languages=("en", "en-US", "en-GB"))
    except NoTranscriptFound:
        try:
            transcript_list = api.list(video_id)
            transcript = next(iter(transcript_list))
            fetched = transcript.fetch()
        except YouTubeTranscriptApiException as exc:
            raise InputValidationError(
                "No usable YouTube transcript was available for that video. If you have "
                "a transcript file, upload it directly."
            ) from exc
        except StopIteration as exc:
            raise InputValidationError(
                "No usable YouTube transcript was available for that video. If you have "
                "a transcript file, upload it directly."
            ) from exc
    except TranscriptsDisabled as exc:
        raise InputValidationError(
            "pdf-it could not find public captions for that YouTube video. Paste or "
            "upload the transcript text directly, or upload audio for transcription."
        ) from exc
    except VideoUnavailable as exc:
        raise InputValidationError(
            "That YouTube video is unavailable. Check the link or video ID and try again."
        ) from exc
    except YouTubeTranscriptApiException as exc:
        raise InputValidationError(
            "pdf-it could not fetch a transcript from YouTube right now. If captions "
            "exist elsewhere, upload the transcript text directly."
        ) from exc

    lines = [
        entry["text"].strip()
        for entry in fetched.to_raw_data()
        if entry.get("text", "").strip()
    ]
    if not lines:
        raise InputValidationError("That YouTube transcript was empty.")
    return "\n".join(lines)


def extract_google_doc_id(url: str) -> str | None:
    """Extract a Google Docs document id from a browser URL."""
    match = re.search(r"/document/d/([a-zA-Z0-9_-]+)", url)
    return match.group(1) if match else None


def extract_youtube_video_id(value: str) -> str | None:
    """Extract a YouTube video id from common URL shapes or a raw 11-character id."""
    cleaned = value.strip()
    if YOUTUBE_VIDEO_ID_PATTERN.fullmatch(cleaned):
        return cleaned

    candidate = cleaned
    lower_candidate = candidate.lower()
    if "://" not in candidate and (
        lower_candidate.startswith(("youtube.com/", "www.youtube.com/", "m.youtube.com/"))
        or lower_candidate.startswith(("music.youtube.com/", "youtu.be/", "www.youtu.be/"))
    ):
        candidate = f"https://{candidate}"

    parsed = urlparse(candidate)
    host = parsed.netloc.lower().split(":", 1)[0]
    video_id: str | None = None

    if host in YOUTUBE_SHORT_HOSTS:
        video_id = next((part for part in parsed.path.split("/") if part), None)
    elif host.endswith(YOUTUBE_HOST_SUFFIXES):
        if parsed.path.rstrip("/") == "/watch":
            query = parse_qs(parsed.query)
            video_id = query.get("v", [None])[0] or query.get("vi", [None])[0]
        else:
            parts = [part for part in parsed.path.split("/") if part]
            if len(parts) >= 2 and parts[0] in {"embed", "shorts", "live", "v"}:
                video_id = parts[1]

    if video_id and YOUTUBE_VIDEO_ID_PATTERN.fullmatch(video_id):
        return video_id
    return None


def normalized_extension(filename: str) -> str:
    parts = filename.rsplit(".", 1)
    return f".{parts[1].lower()}" if len(parts) == 2 else ""


def render_source_segment(label: str, text: str) -> str:
    return f"[Source: {label}]\n{text.strip()}".rstrip()


def truncate_text(text: str, limit: int) -> str:
    if len(text) <= limit:
        return text
    marker = "\n\n[Truncated to fit the shared source limit.]"
    if limit <= len(marker):
        return text[:limit]
    clipped = text[: limit - len(marker)].rstrip()
    last_break = max(clipped.rfind("\n"), clipped.rfind(" "))
    if last_break > 200:
        clipped = clipped[:last_break].rstrip()
    return f"{clipped}{marker}"


def unsupported_extension_message(extension: str) -> str:
    if not extension:
        return (
            "That file has no extension, so pdf-it cannot tell how to read it. "
            "Rename it with a supported extension or copy the contents into the text box."
        )
    if extension in UNSUPPORTED_CODE_EXTENSIONS:
        return (
            f"`{extension}` is not one of the supported code file types. "
            "Copy the file contents into the source box, or save/export it as `.txt` "
            "and upload that instead."
        )
    if extension in VIDEO_EXTENSIONS:
        return (
            f"`{extension}` video files are not supported yet. "
            "Upload a transcript if you have one, or paste the transcript text directly."
        )
    if extension.startswith("."):
        return (
            f"`{extension}` is not a supported upload type yet. "
            "Try converting it to `.txt`, `.md`, `.docx`, `.pdf`, `.csv`, or `.xlsx` "
            "before uploading."
        )
    return "That file type is not supported yet."
